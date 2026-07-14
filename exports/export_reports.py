"""
Capstone Project — Report Exporter
Generates Excel, PDF, and Word reports from the Adventure Works Gold layer.
Falls back to local CSV files if Snowflake is unavailable.

Requirements:
    pip install snowflake-connector-python pandas openpyxl reportlab python-docx matplotlib
"""

import io
import os
import sys
import warnings
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import pandas as pd

warnings.filterwarnings("ignore")
sys.path.insert(0, str(Path(__file__).parent.parent))
from config.snowflake_config import SNOWFLAKE_CONFIG

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE        = Path(__file__).parent.parent
CSV_DIR     = BASE / "capstone_project_dataset"
OUT_DIR     = Path(__file__).parent / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)

TODAY   = datetime.now().strftime("%Y-%m-%d")
XLSX    = OUT_DIR / f"capstone_report_{TODAY}.xlsx"
PDF     = OUT_DIR / f"capstone_report_{TODAY}.pdf"
DOCX    = OUT_DIR / f"capstone_report_{TODAY}.docx"

DB   = "ADVENTURE_WORKS_DB"
GOLD = f"{DB}.GOLD"

# ── Colour Palette ─────────────────────────────────────────────────────────────
NAVY    = "#1F3864"
BLUE    = "#2E75B6"
TEAL    = "#00B0F0"
GREEN   = "#375623"
AMBER   = "#F4B942"
RED     = "#C00000"
GREY    = "#F2F2F2"
WHITE   = "#FFFFFF"
PALETTE = [BLUE, TEAL, AMBER, RED, GREEN, NAVY,
           "#7030A0", "#FF6600", "#00B050", "#FF0000", "#4472C4", "#ED7D31"]


# ==============================================================================
# 1.  DATA LOADING
# ==============================================================================

def _q(conn, sql):
    cur = conn.cursor()
    cur.execute(sql)
    cols = [d[0] for d in cur.description]
    rows = cur.fetchall()
    cur.close()
    return pd.DataFrame(rows, columns=cols)


def load_from_snowflake():
    import snowflake.connector
    cfg = dict(SNOWFLAKE_CONFIG); cfg["schema"] = "GOLD"
    conn = snowflake.connector.connect(**cfg)
    d = {}

    d["sales_by_year"] = _q(conn, f"""
        SELECT c.YEAR,
               COUNT(DISTINCT f.ORDER_NUMBER)          AS ORDERS,
               SUM(f.ORDER_QUANTITY)                   AS UNITS,
               ROUND(SUM(f.GROSS_REVENUE),2)           AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)            AS PROFIT,
               ROUND(AVG(f.PROFIT_MARGIN_PCT),2)       AS MARGIN_PCT
        FROM {GOLD}.FACT_SALES f
        JOIN {GOLD}.DIM_CALENDAR c ON f.CALENDAR_KEY=c.CALENDAR_KEY
        GROUP BY c.YEAR ORDER BY c.YEAR""")

    d["monthly_trend"] = _q(conn, f"""
        SELECT c.YEAR, c.MONTH_NUMBER AS MONTH_NUM, c.MONTH_NAME,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)   AS PROFIT
        FROM {GOLD}.FACT_SALES f
        JOIN {GOLD}.DIM_CALENDAR c ON f.CALENDAR_KEY=c.CALENDAR_KEY
        GROUP BY 1,2,3 ORDER BY 1,2""")

    d["category_revenue"] = _q(conn, f"""
        SELECT p.CATEGORY_NAME,
               SUM(f.ORDER_QUANTITY)           AS UNITS_SOLD,
               ROUND(SUM(f.GROSS_REVENUE),2)   AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)    AS PROFIT,
               ROUND(AVG(f.PROFIT_MARGIN_PCT),2) AS MARGIN_PCT
        FROM {GOLD}.FACT_SALES f
        JOIN {GOLD}.DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        GROUP BY 1 ORDER BY REVENUE DESC""")

    d["top_products"] = _q(conn, f"""
        SELECT p.PRODUCT_NAME, p.CATEGORY_NAME, p.PRICE_TIER,
               SUM(f.ORDER_QUANTITY)           AS UNITS_SOLD,
               ROUND(SUM(f.GROSS_REVENUE),2)   AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)    AS PROFIT
        FROM {GOLD}.FACT_SALES f
        JOIN {GOLD}.DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        GROUP BY 1,2,3 ORDER BY REVENUE DESC LIMIT 15""")

    d["territory"] = _q(conn, f"""
        SELECT t.COUNTRY, t.REGION,
               COUNT(DISTINCT f.ORDER_NUMBER)  AS ORDERS,
               ROUND(SUM(f.GROSS_REVENUE),2)   AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)    AS PROFIT
        FROM {GOLD}.FACT_SALES f
        JOIN {GOLD}.DIM_TERRITORY t ON f.TERRITORY_KEY=t.TERRITORY_KEY
        GROUP BY 1,2 ORDER BY REVENUE DESC""")

    d["customer_segment"] = _q(conn, f"""
        SELECT c.INCOME_BAND, c.OCCUPATION, c.GENDER,
               COUNT(DISTINCT c.CUSTOMER_KEY)                              AS CUSTOMERS,
               ROUND(SUM(f.GROSS_REVENUE),2)                               AS REVENUE,
               ROUND(SUM(f.GROSS_REVENUE)/COUNT(DISTINCT c.CUSTOMER_KEY),2) AS REV_PER_CUST
        FROM {GOLD}.FACT_SALES f
        JOIN {GOLD}.DIM_CUSTOMER c ON f.CUSTOMER_KEY=c.CUSTOMER_KEY
        GROUP BY 1,2,3 ORDER BY REVENUE DESC""")

    d["returns"] = _q(conn, f"""
        SELECT p.CATEGORY_NAME,
               SUM(r.RETURN_QUANTITY)                          AS UNITS_RETURNED,
               COALESCE(SUM(f.ORDER_QUANTITY),0)               AS UNITS_SOLD,
               ROUND(SUM(r.RETURN_QUANTITY)*100.0
                     /NULLIF(SUM(f.ORDER_QUANTITY),0),2)       AS RETURN_RATE_PCT
        FROM {GOLD}.FACT_RETURNS r
        JOIN {GOLD}.DIM_PRODUCT p ON r.PRODUCT_KEY=p.PRODUCT_KEY
        LEFT JOIN {GOLD}.FACT_SALES f ON f.PRODUCT_KEY=r.PRODUCT_KEY
        GROUP BY 1 ORDER BY UNITS_RETURNED DESC""")

    conn.close()
    return d


def load_from_csv():
    MN = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",
          7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}

    def read(fname, **kw):
        return pd.read_csv(CSV_DIR / fname, encoding="latin1", **kw)

    sales_frames = []
    for yr in (2020, 2021, 2022):
        f = read(f"Sales Data {yr}.csv")
        f["Year"] = yr
        sales_frames.append(f)
    sales = pd.concat(sales_frames, ignore_index=True)

    customers = read("Customer Lookup.csv")
    products  = read("Product Lookup.csv")
    cats      = read("Product Categories Lookup.csv")
    subcats   = read("Subcategories Lookup.csv")
    territory = read("Territory Lookup.csv")
    returns   = read("Returns Data.csv")

    for df, col in [(sales,"CustomerKey"),(sales,"ProductKey"),(sales,"TerritoryKey"),
                    (customers,"CustomerKey"),(products,"ProductKey"),
                    (products,"ProductSubcategoryKey"),(subcats,"ProductSubcategoryKey"),
                    (subcats,"ProductCategoryKey"),(cats,"ProductCategoryKey"),
                    (territory,"SalesTerritoryKey"),(returns,"ProductKey"),(returns,"TerritoryKey")]:
        df[col] = pd.to_numeric(df[col], errors="coerce")

    subcats2 = subcats.merge(cats, on="ProductCategoryKey", how="left")
    products2 = products.merge(subcats2, on="ProductSubcategoryKey", how="left")
    products2.columns = [c.strip() for c in products2.columns]
    if "CategoryName_x" in products2.columns:
        products2["CategoryName"] = products2["CategoryName_x"]

    sales = (sales
             .merge(customers[["CustomerKey","AnnualIncome","Gender","Occupation","MaritalStatus"]],
                    on="CustomerKey", how="left")
             .merge(products2[["ProductKey","ProductName","CategoryName","ProductCost","ProductPrice"]],
                    on="ProductKey", how="left")
             .merge(territory[["SalesTerritoryKey","Region","Country"]],
                    left_on="TerritoryKey", right_on="SalesTerritoryKey", how="left"))

    for c in ("ProductPrice","ProductCost","OrderQuantity"):
        sales[c] = pd.to_numeric(sales.get(c, 0), errors="coerce").fillna(0)

    sales["Revenue"] = sales["OrderQuantity"] * sales["ProductPrice"]
    sales["Cost"]    = sales["OrderQuantity"] * sales["ProductCost"]
    sales["Profit"]  = sales["Revenue"] - sales["Cost"]
    sales["Margin"]  = (sales["Profit"] / sales["Revenue"].replace(0, float("nan"))) * 100

    # Parse month from OrderDate (format 2020-01-01)
    sales["OrderDate"] = pd.to_datetime(sales["OrderDate"], errors="coerce")
    sales["Month"]     = sales["OrderDate"].dt.month
    sales["Year"]      = sales["Year"].astype(int)

    def price_tier(p):
        if p < 50:   return "Budget (<$50)"
        if p < 200:  return "Mid ($50-200)"
        if p < 500:  return "Premium ($200-500)"
        return "Luxury (>$500)"
    sales["PriceTier"] = sales["ProductPrice"].apply(price_tier)

    d = {}

    sy = (sales.groupby("Year")
               .agg(ORDERS=("OrderNumber","nunique"),
                    UNITS=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"),
                    PROFIT=("Profit","sum"))
               .reset_index().rename(columns={"Year":"YEAR"}))
    sy["MARGIN_PCT"] = (sy["PROFIT"]/sy["REVENUE"]*100).round(2)
    sy[["REVENUE","PROFIT"]] = sy[["REVENUE","PROFIT"]].round(2)
    d["sales_by_year"] = sy

    mt = (sales.groupby(["Year","Month"])["Revenue"].sum().reset_index()
               .rename(columns={"Year":"YEAR","Month":"MONTH_NUM","Revenue":"REVENUE"}))
    mt["MONTH_NAME"] = mt["MONTH_NUM"].map(MN)
    mt["REVENUE"] = mt["REVENUE"].round(2)
    d["monthly_trend"] = mt[["YEAR","MONTH_NUM","MONTH_NAME","REVENUE"]]

    cr = (sales.groupby("CategoryName")
               .agg(UNITS_SOLD=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"),
                    PROFIT=("Profit","sum"),
                    MARGIN_PCT=("Margin","mean"))
               .reset_index().rename(columns={"CategoryName":"CATEGORY_NAME"})
               .sort_values("REVENUE", ascending=False))
    cr[["REVENUE","PROFIT","MARGIN_PCT"]] = cr[["REVENUE","PROFIT","MARGIN_PCT"]].round(2)
    d["category_revenue"] = cr

    tp = (sales.groupby(["ProductName","CategoryName","PriceTier"])
               .agg(UNITS_SOLD=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"),
                    PROFIT=("Profit","sum"))
               .reset_index()
               .rename(columns={"ProductName":"PRODUCT_NAME","CategoryName":"CATEGORY_NAME",
                                 "PriceTier":"PRICE_TIER"})
               .sort_values("REVENUE", ascending=False).head(15))
    tp[["REVENUE","PROFIT"]] = tp[["REVENUE","PROFIT"]].round(2)
    d["top_products"] = tp

    terr = (sales.groupby(["Country","Region"])
                 .agg(ORDERS=("OrderNumber","nunique"), REVENUE=("Revenue","sum"), PROFIT=("Profit","sum"))
                 .reset_index().rename(columns={"Country":"COUNTRY","Region":"REGION"})
                 .sort_values("REVENUE", ascending=False))
    terr[["REVENUE","PROFIT"]] = terr[["REVENUE","PROFIT"]].round(2)
    d["territory"] = terr

    seg = (sales.groupby(["Occupation","Gender"])
                .agg(CUSTOMERS=("CustomerKey","nunique"),
                     REVENUE=("Revenue","sum"))
                .reset_index()
                .rename(columns={"Occupation":"OCCUPATION","Gender":"GENDER"})
                .sort_values("REVENUE", ascending=False))
    seg["REVENUE"] = seg["REVENUE"].round(2)
    seg["REV_PER_CUST"] = (seg["REVENUE"]/seg["CUSTOMERS"]).round(2)
    d["customer_segment"] = seg

    ret_p = (returns.groupby("ProductKey")["ReturnQuantity"].sum()
                    .reset_index().rename(columns={"ReturnQuantity":"UNITS_RETURNED"}))
    sold_p = (sales.groupby("ProductKey")
                   .agg(UNITS_SOLD=("OrderQuantity","sum"),
                        CATEGORY_NAME=("CategoryName","first"))
                   .reset_index())
    rc = (sold_p.merge(ret_p, on="ProductKey", how="left"))
    rc["UNITS_RETURNED"] = rc["UNITS_RETURNED"].fillna(0).astype(int)
    rc = (rc.groupby("CATEGORY_NAME")
            .agg(UNITS_SOLD=("UNITS_SOLD","sum"),
                 UNITS_RETURNED=("UNITS_RETURNED","sum"))
            .reset_index())
    rc["RETURN_RATE_PCT"] = (rc["UNITS_RETURNED"]/rc["UNITS_SOLD"]*100).round(2)
    d["returns"] = rc.sort_values("UNITS_RETURNED", ascending=False)

    return d


def get_data():
    try:
        import snowflake.connector
        d = load_from_snowflake()
        print("[OK] Loaded from Snowflake")
        return d
    except Exception as e:
        print(f"[INFO] Snowflake unavailable ({type(e).__name__}: {e}) — using CSV fallback")
    try:
        d = load_from_csv()
        print("[OK] Loaded from CSV")
        return d
    except Exception as e2:
        raise RuntimeError(f"Both Snowflake and CSV loading failed: {e2}") from e2


# ==============================================================================
# 2.  KPIs
# ==============================================================================

def compute_kpis(d):
    sy = d["sales_by_year"]
    cr = d["category_revenue"]
    terr = d["territory"]
    return {
        "total_revenue": sy["REVENUE"].sum(),
        "total_profit":  sy["PROFIT"].sum(),
        "total_orders":  sy["ORDERS"].sum(),
        "avg_margin":    (sy["PROFIT"].sum()/sy["REVENUE"].sum()*100),
        "top_category":  cr.iloc[0]["CATEGORY_NAME"] if len(cr) else "N/A",
        "top_country":   terr.iloc[0]["COUNTRY"]     if len(terr) else "N/A",
        "years":         sorted(sy["YEAR"].tolist()),
    }


def fmt(v):
    if v >= 1_000_000: return f"${v/1_000_000:.1f}M"
    if v >= 1_000:     return f"${v/1_000:.0f}K"
    return f"${v:.0f}"


# ==============================================================================
# 3.  CHARTS  (return bytes)
# ==============================================================================

def _save(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def chart_revenue_bar(d):
    sy = d["sales_by_year"]
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(sy["YEAR"].astype(str), sy["REVENUE"]/1e6, color=PALETTE[:len(sy)], edgecolor="white")
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.05,
                f"${b.get_height():.1f}M", ha="center", va="bottom", fontsize=9, fontweight="bold")
    ax.set_title("Annual Revenue ($M)", fontsize=12, fontweight="bold", color=NAVY)
    ax.set_ylabel("Revenue ($M)"); ax.set_xlabel("Year")
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _save(fig)


def chart_monthly_trend(d):
    mt = d["monthly_trend"]
    fig, ax = plt.subplots(figsize=(9, 3.5))
    for i, yr in enumerate(sorted(mt["YEAR"].unique())):
        sub = mt[mt["YEAR"]==yr].sort_values("MONTH_NUM")
        ax.plot(sub["MONTH_NUM"], sub["REVENUE"]/1e6, marker="o",
                label=str(yr), color=PALETTE[i], linewidth=2)
    ax.set_title("Monthly Revenue Trend by Year ($M)", fontsize=12, fontweight="bold", color=NAVY)
    ax.set_xlabel("Month"); ax.set_ylabel("Revenue ($M)")
    ax.set_xticks(range(1,13))
    ax.set_xticklabels(["J","F","M","A","M","J","J","A","S","O","N","D"])
    ax.legend(frameon=False)
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _save(fig)


def chart_category_bar(d):
    cr = d["category_revenue"].head(8)
    fig, ax = plt.subplots(figsize=(7, 3.5))
    colors_bar = PALETTE[:len(cr)]
    bars = ax.barh(cr["CATEGORY_NAME"], cr["REVENUE"]/1e6, color=colors_bar)
    for b in bars:
        ax.text(b.get_width()+0.02, b.get_y()+b.get_height()/2,
                f"${b.get_width():.1f}M", va="center", fontsize=8)
    ax.set_title("Revenue by Category ($M)", fontsize=12, fontweight="bold", color=NAVY)
    ax.set_xlabel("Revenue ($M)")
    ax.invert_yaxis()
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _save(fig)


def chart_category_pie(d):
    cr = d["category_revenue"].head(6)
    fig, ax = plt.subplots(figsize=(5.5, 4))
    wedges, texts, autotexts = ax.pie(
        cr["REVENUE"], labels=None, autopct="%1.1f%%",
        colors=PALETTE[:len(cr)], startangle=140,
        pctdistance=0.8, wedgeprops={"linewidth":1,"edgecolor":"white"})
    for at in autotexts:
        at.set_fontsize(8)
    ax.legend(wedges, cr["CATEGORY_NAME"], loc="lower center",
              bbox_to_anchor=(0.5,-0.15), ncol=2, fontsize=8, frameon=False)
    ax.set_title("Revenue Share by Category", fontsize=12, fontweight="bold", color=NAVY)
    fig.tight_layout()
    return _save(fig)


def chart_territory_bar(d):
    terr = d["territory"].head(8)
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(terr["COUNTRY"], terr["REVENUE"]/1e6,
                  color=PALETTE[:len(terr)], edgecolor="white")
    plt.xticks(rotation=30, ha="right", fontsize=8)
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.05,
                f"${b.get_height():.1f}M", ha="center", fontsize=7)
    ax.set_title("Revenue by Country ($M)", fontsize=12, fontweight="bold", color=NAVY)
    ax.set_ylabel("Revenue ($M)")
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _save(fig)


def chart_occupation_bar(d):
    seg = (d["customer_segment"].groupby("OCCUPATION")["REVENUE"]
               .sum().reset_index().sort_values("REVENUE", ascending=False))
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(seg["OCCUPATION"], seg["REVENUE"]/1e6,
                  color=PALETTE[:len(seg)], edgecolor="white")
    plt.xticks(rotation=20, ha="right", fontsize=8)
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.02,
                f"${b.get_height():.1f}M", ha="center", fontsize=7)
    ax.set_title("Revenue by Occupation ($M)", fontsize=12, fontweight="bold", color=NAVY)
    ax.set_ylabel("Revenue ($M)")
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _save(fig)


def chart_gender_pie(d):
    seg = d["customer_segment"]
    gs = seg.groupby("GENDER")["REVENUE"].sum().reset_index()
    fig, ax = plt.subplots(figsize=(5, 4))
    wedges, texts, autotexts = ax.pie(
        gs["REVENUE"], labels=gs["GENDER"], autopct="%1.1f%%",
        colors=[BLUE, AMBER], startangle=90,
        wedgeprops={"linewidth":2,"edgecolor":"white"})
    ax.set_title("Revenue by Gender", fontsize=12, fontweight="bold", color=NAVY)
    fig.tight_layout()
    return _save(fig)


def chart_returns_bar(d):
    rc = d["returns"].head(8)
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(rc["CATEGORY_NAME"], rc["UNITS_RETURNED"],
                  color=[RED]*len(rc), edgecolor="white", alpha=0.85)
    plt.xticks(rotation=20, ha="right", fontsize=8)
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.5,
                str(int(b.get_height())), ha="center", fontsize=8)
    ax.set_title("Units Returned by Category", fontsize=12, fontweight="bold", color=RED)
    ax.set_ylabel("Units Returned")
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _save(fig)


def make_all_charts(d):
    print("  Generating charts...")
    return {
        "revenue_bar":    chart_revenue_bar(d),
        "monthly_trend":  chart_monthly_trend(d),
        "category_bar":   chart_category_bar(d),
        "category_pie":   chart_category_pie(d),
        "territory_bar":  chart_territory_bar(d),
        "occupation_bar": chart_occupation_bar(d),
        "gender_pie":     chart_gender_pie(d),
        "returns_bar":    chart_returns_bar(d),
    }


# ==============================================================================
# 4.  EXCEL REPORT
# ==============================================================================

def build_excel(d, charts, kpis):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.drawing.image import Image as XLImage

    wb = Workbook()
    wb.remove(wb.active)

    navy_fill  = PatternFill("solid", fgColor="1F3864")
    blue_fill  = PatternFill("solid", fgColor="2E75B6")
    grey_fill  = PatternFill("solid", fgColor="F2F2F2")
    white_font = Font(color="FFFFFF", bold=True, size=11)
    hdr_font   = Font(color="FFFFFF", bold=True, size=10)
    title_font = Font(color="1F3864", bold=True, size=14)
    bold_font  = Font(bold=True)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def xl_img(ws, chart_bytes, anchor, w=500, h=270):
        img = XLImage(io.BytesIO(chart_bytes))
        img.width = w; img.height = h
        ws.add_image(img, anchor)

    def xl_title(ws, text, row, span, fill=navy_fill):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)
        c = ws.cell(row=row, column=1, value=text)
        c.font = white_font; c.fill = fill; c.alignment = Alignment(horizontal="center")

    def xl_table(ws, df, start_row, hdr_fill=blue_fill):
        for ci, col in enumerate(df.columns, 1):
            c = ws.cell(row=start_row, column=ci, value=col)
            c.font = hdr_font; c.fill = hdr_fill; c.border = border
            c.alignment = Alignment(horizontal="center")
        for ri, (_, row) in enumerate(df.iterrows(), start_row+1):
            f = grey_fill if ri % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
            for ci, val in enumerate(row, 1):
                c = ws.cell(row=ri, column=ci, value=val)
                c.border = border; c.fill = f
                if isinstance(val, float):
                    c.number_format = "#,##0.00"

    # ── Sheet 1: Executive Summary ──────────────────────────────────────────────
    ws1 = wb.create_sheet("Executive Summary")
    ws1.sheet_view.showGridLines = False
    ws1.column_dimensions["A"].width = 3

    ws1.merge_cells("B2:H2")
    c = ws1["B2"]; c.value = "ADVENTURE WORKS — CAPSTONE ANALYTICS REPORT"
    c.font = Font(color="1F3864", bold=True, size=16); c.alignment = Alignment(horizontal="center")
    ws1.merge_cells("B3:H3")
    c = ws1["B3"]; c.value = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Account: uq57089.ap-southeast-7.aws"
    c.font = Font(italic=True, size=10, color="666666"); c.alignment = Alignment(horizontal="center")

    kpi_data = [
        ("Total Revenue",  fmt(kpis["total_revenue"])),
        ("Total Profit",   fmt(kpis["total_profit"])),
        ("Total Orders",   f"{int(kpis['total_orders']):,}"),
        ("Avg Margin",     f"{kpis['avg_margin']:.1f}%"),
        ("Top Category",   kpis["top_category"]),
        ("Top Country",    kpis["top_country"]),
    ]
    for i, (label, value) in enumerate(kpi_data):
        col = 2 + (i * 2)
        ws1.merge_cells(start_row=5, start_column=col, end_row=5, end_column=col+1)
        ws1.merge_cells(start_row=6, start_column=col, end_row=6, end_column=col+1)
        lc = ws1.cell(row=5, column=col, value=label)
        lc.font = Font(bold=True, color="FFFFFF", size=10)
        lc.fill = navy_fill; lc.alignment = Alignment(horizontal="center")
        vc = ws1.cell(row=6, column=col, value=value)
        vc.font = Font(bold=True, size=12, color="1F3864")
        vc.fill = grey_fill; vc.alignment = Alignment(horizontal="center")

    xl_img(ws1, charts["revenue_bar"],   "B8",  w=450, h=260)
    xl_img(ws1, charts["monthly_trend"], "K8",  w=550, h=260)

    # ── Sheet 2: Sales Performance ──────────────────────────────────────────────
    ws2 = wb.create_sheet("Sales Performance")
    ws2.sheet_view.showGridLines = False
    xl_title(ws2, "Annual Sales Summary", 1, 6)
    xl_table(ws2, d["sales_by_year"], start_row=2)
    xl_img(ws2, charts["revenue_bar"],   "A8",  w=480, h=280)
    xl_img(ws2, charts["monthly_trend"], "J8",  w=540, h=280)
    xl_title(ws2, "Monthly Revenue Trend", 24, 4)
    xl_table(ws2, d["monthly_trend"].head(36), start_row=25)

    # ── Sheet 3: Product Analysis ───────────────────────────────────────────────
    ws3 = wb.create_sheet("Product Analysis")
    ws3.sheet_view.showGridLines = False
    xl_title(ws3, "Revenue by Category", 1, 5)
    xl_table(ws3, d["category_revenue"], start_row=2)
    xl_img(ws3, charts["category_bar"], "A10", w=480, h=280)
    xl_img(ws3, charts["category_pie"], "J10", w=420, h=300)
    xl_title(ws3, "Top 15 Products by Revenue", 26, 6)
    xl_table(ws3, d["top_products"], start_row=27)

    # ── Sheet 4: Territory Analysis ─────────────────────────────────────────────
    ws4 = wb.create_sheet("Territory Analysis")
    ws4.sheet_view.showGridLines = False
    xl_title(ws4, "Revenue by Territory", 1, 5)
    xl_table(ws4, d["territory"], start_row=2)
    xl_img(ws4, charts["territory_bar"], "A14", w=540, h=300)

    # ── Sheet 5: Customer Segmentation ──────────────────────────────────────────
    ws5 = wb.create_sheet("Customer Segmentation")
    ws5.sheet_view.showGridLines = False
    xl_title(ws5, "Customer Segment Revenue", 1, 6)
    xl_table(ws5, d["customer_segment"].head(20), start_row=2)
    xl_img(ws5, charts["occupation_bar"], "A14", w=500, h=280)
    xl_img(ws5, charts["gender_pie"],     "J14", w=400, h=300)

    # ── Sheet 6: Returns Analysis ────────────────────────────────────────────────
    ws6 = wb.create_sheet("Returns Analysis")
    ws6.sheet_view.showGridLines = False
    xl_title(ws6, "Returns by Category", 1, 5, fill=PatternFill("solid", fgColor="C00000"))
    xl_table(ws6, d["returns"], start_row=2, hdr_fill=PatternFill("solid", fgColor="C00000"))
    xl_img(ws6, charts["returns_bar"], "A10", w=520, h=280)

    wb.save(XLSX)
    print(f"  Excel saved -> {XLSX}")


# ==============================================================================
# 5.  PDF REPORT
# ==============================================================================

def build_pdf(d, charts, kpis):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer, PageBreak, HRFlowable, Image as RLImage)

    LNAV = rl_colors.HexColor(NAVY)
    LBLU = rl_colors.HexColor(BLUE)
    LRED = rl_colors.HexColor(RED)

    doc  = SimpleDocTemplate(str(PDF), pagesize=landscape(A4),
                              rightMargin=1.5*cm, leftMargin=1.5*cm,
                              topMargin=2*cm,    bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    H0 = ParagraphStyle("h0", parent=styles["Title"],   textColor=LNAV, fontSize=20, spaceAfter=6)
    H1 = ParagraphStyle("h1", parent=styles["Heading1"],textColor=LNAV, fontSize=14, spaceAfter=4)
    H2 = ParagraphStyle("h2", parent=styles["Heading2"],textColor=LBLU, fontSize=11, spaceAfter=3)
    SM = ParagraphStyle("sm", parent=styles["Normal"],  fontSize=9,  spaceAfter=3, leading=13)

    def rl_img(key, width=12*cm, height=7*cm):
        return RLImage(io.BytesIO(charts[key]), width=width, height=height)

    def pdf_table(df, hdr_color=LBLU):
        data = [list(df.columns)] + [[str(v) for v in r] for _, r in df.iterrows()]
        tbl  = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",     (0,0),(-1,0),  hdr_color),
            ("TEXTCOLOR",      (0,0),(-1,0),  rl_colors.white),
            ("FONTNAME",       (0,0),(-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",       (0,0),(-1,-1), 7.5),
            ("ROWBACKGROUNDS", (0,1),(-1,-1), [rl_colors.HexColor("#EBF3FB"), rl_colors.white]),
            ("GRID",           (0,0),(-1,-1), 0.3, rl_colors.grey),
            ("VALIGN",         (0,0),(-1,-1), "MIDDLE"),
            ("ALIGN",          (0,0),(-1,-1), "CENTER"),
        ]))
        return tbl

    PAGE_W = landscape(A4)[0] - 3*cm
    half = PAGE_W / 2

    story = []

    # Cover
    story += [
        Paragraph("Adventure Works", H0),
        Paragraph("Snowflake Capstone Analytics Report", H1),
        Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} &nbsp;|&nbsp; "
                  f"Account: uq57089.ap-southeast-7.aws &nbsp;|&nbsp; DB: ADVENTURE_WORKS_DB", SM),
        Spacer(1, 0.3*cm),
        HRFlowable(width="100%", thickness=2, color=LNAV),
        Spacer(1, 0.5*cm),
    ]
    kpi_rows = [
        ["Metric", "Value"],
        ["Total Revenue",  fmt(kpis["total_revenue"])],
        ["Total Profit",   fmt(kpis["total_profit"])],
        ["Total Orders",   f"{int(kpis['total_orders']):,}"],
        ["Avg Margin %",   f"{kpis['avg_margin']:.1f}%"],
        ["Top Category",   kpis["top_category"]],
        ["Top Country",    kpis["top_country"]],
    ]
    kpi_tbl = Table(kpi_rows, colWidths=[7*cm, 7*cm])
    kpi_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,0), LNAV),
        ("TEXTCOLOR",  (0,0),(-1,0), rl_colors.white),
        ("FONTNAME",   (0,0),(-1,-1),"Helvetica-Bold"),
        ("FONTSIZE",   (0,0),(-1,-1), 11),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.HexColor("#EBF3FB"),rl_colors.white]),
        ("GRID",       (0,0),(-1,-1), 0.5, rl_colors.grey),
        ("ALIGN",      (0,0),(-1,-1),"CENTER"),
    ]))
    story += [kpi_tbl, PageBreak()]

    # Section 1 — Sales
    story += [Paragraph("1. Sales Performance", H1),
              pdf_table(d["sales_by_year"]), Spacer(1,0.4*cm)]
    img_row = Table([[rl_img("revenue_bar",  half-0.5*cm, 8*cm),
                      rl_img("monthly_trend",half+0.5*cm, 8*cm)]],
                    colWidths=[half-0.5*cm, half+0.5*cm])
    story += [img_row, PageBreak()]

    # Section 2 — Products
    story += [Paragraph("2. Product Analysis", H1),
              pdf_table(d["category_revenue"]), Spacer(1,0.4*cm)]
    img_row2 = Table([[rl_img("category_bar", half-0.5*cm, 8*cm),
                       rl_img("category_pie", half+0.5*cm, 8*cm)]],
                     colWidths=[half-0.5*cm, half+0.5*cm])
    story += [img_row2, Spacer(1,0.4*cm),
              Paragraph("Top 15 Products by Revenue", H2),
              pdf_table(d["top_products"]), PageBreak()]

    # Section 3 — Territory
    story += [Paragraph("3. Territory Analysis", H1),
              pdf_table(d["territory"]), Spacer(1,0.4*cm),
              rl_img("territory_bar", PAGE_W*0.7, 9*cm), PageBreak()]

    # Section 4 — Customers
    story += [Paragraph("4. Customer Segmentation", H1),
              pdf_table(d["customer_segment"].head(15)), Spacer(1,0.4*cm)]
    img_row3 = Table([[rl_img("occupation_bar", half-0.5*cm, 8*cm),
                       rl_img("gender_pie",     half+0.5*cm, 8*cm)]],
                     colWidths=[half-0.5*cm, half+0.5*cm])
    story += [img_row3, PageBreak()]

    # Section 5 — Returns
    story += [Paragraph("5. Returns Analysis", H1),
              pdf_table(d["returns"], hdr_color=LRED), Spacer(1,0.4*cm),
              rl_img("returns_bar", PAGE_W*0.65, 8*cm)]

    doc.build(story)
    print(f"  PDF  saved -> {PDF}")


# ==============================================================================
# 6.  WORD REPORT
# ==============================================================================

def build_word(d, charts, kpis):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY_RGB  = RGBColor(0x1F, 0x4E, 0x79)
    BLUE_RGB  = RGBColor(0x2E, 0x75, 0xB6)
    WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
    RED_RGB   = RGBColor(0xC0, 0x00, 0x00)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    def heading(text, level=1, color=NAVY_RGB):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color

    def subheading(text):
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.color.rgb = BLUE_RGB

    def para(text, bold=False, italic=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        return p

    def add_chart(key, width=Inches(5.5)):
        doc.add_picture(io.BytesIO(charts[key]), width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def cell_shade(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def add_table(df, hdr_hex="1F3864"):
        tbl = doc.add_table(rows=1, cols=len(df.columns))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold = True; run.font.color.rgb = WHITE_RGB
            run.font.size = Pt(9)
            cell_shade(hdr[i], hdr_hex)
        for _, row in df.iterrows():
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        return tbl

    # ── Cover Page ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Adventure Works")
    run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = NAVY_RGB

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Snowflake Capstone Analytics Report")
    run2.font.size = Pt(18); run2.font.color.rgb = BLUE_RGB

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run3.font.size = Pt(11); run3.font.italic = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Account: uq57089.ap-southeast-7.aws  |  Database: ADVENTURE_WORKS_DB")
    run4.font.size = Pt(10)

    doc.add_paragraph()

    # KPI table
    heading("Executive Summary KPIs", level=2)
    kpi_rows = [
        ("Total Revenue",  fmt(kpis["total_revenue"])),
        ("Total Profit",   fmt(kpis["total_profit"])),
        ("Total Orders",   f"{int(kpis['total_orders']):,}"),
        ("Avg Margin %",   f"{kpis['avg_margin']:.1f}%"),
        ("Top Category",   kpis["top_category"]),
        ("Top Country",    kpis["top_country"]),
    ]
    kpi_tbl = doc.add_table(rows=1, cols=2)
    kpi_tbl.style = "Table Grid"
    kpi_tbl.rows[0].cells[0].text = "KPI"
    kpi_tbl.rows[0].cells[1].text = "Value"
    for cell in kpi_tbl.rows[0].cells:
        cell_shade(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        cell.paragraphs[0].runs[0].font.bold = True
    for label, value in kpi_rows:
        row = kpi_tbl.add_row().cells
        row[0].text = label; row[1].text = value

    doc.add_page_break()

    # ── Section 1 — Sales ──────────────────────────────────────────────────────
    heading("1. Sales Performance")
    para(f"Total revenue across all years: {fmt(kpis['total_revenue'])} | "
         f"Avg profit margin: {kpis['avg_margin']:.1f}%", size=10)
    add_table(d["sales_by_year"])
    doc.add_paragraph()
    add_chart("revenue_bar",   width=Inches(5.5))
    add_chart("monthly_trend", width=Inches(5.5))
    doc.add_page_break()

    # ── Section 2 — Products ──────────────────────────────────────────────────
    heading("2. Product Analysis")
    subheading("Revenue by Category")
    add_table(d["category_revenue"])
    doc.add_paragraph()
    add_chart("category_bar", width=Inches(5.0))
    add_chart("category_pie", width=Inches(4.5))
    subheading("Top 15 Products by Revenue")
    add_table(d["top_products"])
    doc.add_page_break()

    # ── Section 3 — Territory ─────────────────────────────────────────────────
    heading("3. Territory Analysis")
    add_table(d["territory"])
    doc.add_paragraph()
    add_chart("territory_bar", width=Inches(5.5))
    doc.add_page_break()

    # ── Section 4 — Customers ─────────────────────────────────────────────────
    heading("4. Customer Segmentation")
    add_table(d["customer_segment"].head(15))
    doc.add_paragraph()
    add_chart("occupation_bar", width=Inches(5.0))
    add_chart("gender_pie",     width=Inches(4.0))
    doc.add_page_break()

    # ── Section 5 — Returns ───────────────────────────────────────────────────
    heading("5. Returns Analysis", color=RED_RGB)
    add_table(d["returns"], hdr_hex="C00000")
    doc.add_paragraph()
    add_chart("returns_bar", width=Inches(5.5))

    doc.save(DOCX)
    print(f"  Word saved -> {DOCX}")


# ==============================================================================
# 7.  MAIN
# ==============================================================================

def main():
    print("=" * 55)
    print("  ADVENTURE WORKS — CAPSTONE REPORT GENERATOR")
    print(f"  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 55)

    print("\n[1/4] Loading data...")
    d    = get_data()
    kpis = compute_kpis(d)
    print(f"  Revenue: {fmt(kpis['total_revenue'])}  |  "
          f"Orders: {int(kpis['total_orders']):,}  |  "
          f"Margin: {kpis['avg_margin']:.1f}%")

    print("\n[2/4] Generating charts...")
    charts = make_all_charts(d)
    print(f"  {len(charts)} charts generated")

    print("\n[3/4] Building reports...")
    build_excel(d, charts, kpis)
    build_pdf(d, charts, kpis)
    build_word(d, charts, kpis)

    print("\n[4/4] Complete!")
    print(f"  Excel -> {XLSX}")
    print(f"  PDF   -> {PDF}")
    print(f"  Word  -> {DOCX}")
    print("=" * 55)


if __name__ == "__main__":
    main()

"""
Adventure Works -- Enhanced Report Generator v2
Bar charts, pie charts (donut), and deep analysis across Excel / PDF / Word.

Data  : Snowflake ADVENTURE_WORKS_DB.GOLD  (auto-falls back to local CSVs)
Output: reports/output/adventure_works_report_<date>.[xlsx|pdf|docx]
Run   : python generate_report.py
"""

import io, os, datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import pandas as pd

# ── paths ──────────────────────────────────────────────────────────────────────
BASE  = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CSV   = os.path.join(BASE, "capstone_project_dataset")
OUT   = os.path.join(BASE, "reports", "output")
TODAY = datetime.date.today().strftime("%Y-%m-%d")
XLSX  = os.path.join(OUT, f"adventure_works_report_{TODAY}.xlsx")
PDF   = os.path.join(OUT, f"adventure_works_report_{TODAY}.pdf")
DOCX  = os.path.join(OUT, f"adventure_works_report_{TODAY}.docx")

PALETTE = ["#1F4E79","#2E75B6","#9DC3E6","#BDD7EE",
           "#145A32","#27AE60","#6C3483","#A569BD",
           "#784212","#E59866","#922B21","#F1948A"]

plt.rcParams.update({
    "figure.facecolor":"white", "axes.facecolor":"white",
    "axes.grid":True, "grid.color":"#E0E0E0", "grid.linewidth":0.5,
    "font.size":9, "axes.titlesize":11, "axes.titleweight":"bold",
    "axes.labelsize":9, "xtick.labelsize":8, "ytick.labelsize":8,
})


# ══════════════════════════════════════════════════════════════════════════════
# 1.  HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _query(conn, sql):
    cur = conn.cursor()
    cur.execute(sql)
    return pd.DataFrame(cur.fetchall(), columns=[d[0] for d in cur.description])

def _fmt(v):
    if v >= 1_000_000: return f"${v/1_000_000:.1f}M"
    if v >= 1_000:     return f"${v/1_000:.0f}K"
    return f"${v:.0f}"

def _savefig(fig):
    """Return chart as raw bytes so each format can open a fresh BytesIO."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()   # bytes — never gets closed


# ══════════════════════════════════════════════════════════════════════════════
# 2.  DATA LOADING
# ══════════════════════════════════════════════════════════════════════════════

def load_from_snowflake():
    from connect import get_conn
    conn = get_conn()
    print("  Connected to Snowflake [OK]")
    d = {}

    d["sales_by_year"] = _query(conn, """
        SELECT d.YEAR,
               COUNT(DISTINCT f.ORDER_NUMBER) AS TOTAL_ORDERS,
               SUM(f.ORDER_QUANTITY)          AS UNITS_SOLD,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)   AS PROFIT,
               ROUND(AVG(f.MARGIN_PCT),2)     AS MARGIN_PCT
        FROM FACT_SALES f JOIN DIM_DATE d ON f.ORDER_DATE_KEY=d.DATE_KEY
        GROUP BY d.YEAR ORDER BY d.YEAR""")
    d["sales_by_year"]["YOY_GROWTH_PCT"] = (
        d["sales_by_year"]["REVENUE"].pct_change() * 100).round(2)

    d["monthly_trend"] = _query(conn, """
        SELECT d.YEAR, d.MONTH_NUM, d.MONTH_NAME,
               ROUND(SUM(f.GROSS_REVENUE),2) AS REVENUE
        FROM FACT_SALES f JOIN DIM_DATE d ON f.ORDER_DATE_KEY=d.DATE_KEY
        GROUP BY d.YEAR,d.MONTH_NUM,d.MONTH_NAME ORDER BY d.YEAR,d.MONTH_NUM""")

    d["category_revenue"] = _query(conn, """
        SELECT p.CATEGORY_NAME,
               SUM(f.ORDER_QUANTITY)          AS UNITS_SOLD,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)   AS PROFIT,
               ROUND(AVG(f.MARGIN_PCT),2)     AS MARGIN_PCT
        FROM FACT_SALES f JOIN DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        GROUP BY p.CATEGORY_NAME ORDER BY REVENUE DESC""")

    d["top_products"] = _query(conn, """
        SELECT p.PRODUCT_NAME, p.CATEGORY_NAME,
               SUM(f.ORDER_QUANTITY)          AS UNITS_SOLD,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE,
               ROUND(SUM(f.GROSS_PROFIT),2)   AS PROFIT
        FROM FACT_SALES f JOIN DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        GROUP BY 1,2 ORDER BY REVENUE DESC LIMIT 10""")

    d["territory"] = _query(conn, """
        SELECT t.COUNTRY, t.REGION,
               COUNT(DISTINCT f.ORDER_NUMBER)           AS ORDERS,
               ROUND(SUM(f.GROSS_REVENUE),2)            AS REVENUE,
               ROUND(100*SUM(f.GROSS_REVENUE)
                     /SUM(SUM(f.GROSS_REVENUE)) OVER(),2) AS REVENUE_PCT
        FROM FACT_SALES f JOIN DIM_TERRITORY t ON f.TERRITORY_KEY=t.TERRITORY_KEY
        GROUP BY 1,2 ORDER BY REVENUE DESC""")

    d["gender_split"] = _query(conn, """
        SELECT c.GENDER,
               COUNT(DISTINCT c.CUSTOMER_KEY) AS CUSTOMERS,
               COUNT(DISTINCT f.ORDER_NUMBER) AS ORDERS,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE
        FROM FACT_SALES f JOIN DIM_CUSTOMER c ON f.CUSTOMER_KEY=c.CUSTOMER_KEY
        GROUP BY c.GENDER""")

    d["occupation_revenue"] = _query(conn, """
        SELECT c.OCCUPATION,
               COUNT(DISTINCT c.CUSTOMER_KEY) AS CUSTOMERS,
               COUNT(DISTINCT f.ORDER_NUMBER) AS ORDERS,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE
        FROM FACT_SALES f JOIN DIM_CUSTOMER c ON f.CUSTOMER_KEY=c.CUSTOMER_KEY
        GROUP BY c.OCCUPATION ORDER BY REVENUE DESC""")

    d["returns_by_category"] = _query(conn, """
        SELECT p.CATEGORY_NAME,
               SUM(f.ORDER_QUANTITY)                             AS UNITS_SOLD,
               COALESCE(SUM(r.RETURN_QUANTITY),0)               AS UNITS_RETURNED,
               ROUND(COALESCE(SUM(r.RETURN_QUANTITY),0)
                     /NULLIF(SUM(f.ORDER_QUANTITY),0)*100,2)    AS RETURN_RATE_PCT
        FROM FACT_SALES f JOIN DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        LEFT JOIN FACT_RETURNS r ON f.PRODUCT_KEY=r.PRODUCT_KEY
        GROUP BY 1 ORDER BY RETURN_RATE_PCT DESC""")

    d["price_tier"] = _query(conn, """
        SELECT CASE
                 WHEN p.PRODUCT_PRICE < 50   THEN 'Budget (<$50)'
                 WHEN p.PRODUCT_PRICE < 200  THEN 'Mid-Range ($50-199)'
                 WHEN p.PRODUCT_PRICE < 500  THEN 'Premium ($200-499)'
                 ELSE 'Luxury ($500+)'
               END AS PRICE_TIER,
               COUNT(DISTINCT f.ORDER_NUMBER) AS ORDERS,
               SUM(f.ORDER_QUANTITY)          AS UNITS_SOLD,
               ROUND(SUM(f.GROSS_REVENUE),2)  AS REVENUE,
               ROUND(100*SUM(f.GROSS_REVENUE)
                     /SUM(SUM(f.GROSS_REVENUE)) OVER(),2) AS REVENUE_PCT
        FROM FACT_SALES f JOIN DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        GROUP BY 1 ORDER BY REVENUE DESC""")

    d["customers_detail"] = _query(conn, """
        SELECT c.INCOME_BAND, c.OCCUPATION, c.GENDER, c.MARITAL_STATUS,
               COUNT(DISTINCT c.CUSTOMER_KEY)                          AS CUSTOMERS,
               COUNT(DISTINCT f.ORDER_NUMBER)                          AS ORDERS,
               ROUND(SUM(f.GROSS_REVENUE),2)                           AS REVENUE,
               ROUND(SUM(f.GROSS_REVENUE)/COUNT(DISTINCT c.CUSTOMER_KEY),2) AS REVENUE_PER_CUSTOMER
        FROM FACT_SALES f JOIN DIM_CUSTOMER c ON f.CUSTOMER_KEY=c.CUSTOMER_KEY
        GROUP BY 1,2,3,4 ORDER BY REVENUE DESC""")

    d["returns_detail"] = _query(conn, """
        SELECT p.PRODUCT_NAME, p.CATEGORY_NAME,
               SUM(f.ORDER_QUANTITY)                             AS UNITS_SOLD,
               COALESCE(SUM(r.RETURN_QUANTITY),0)               AS UNITS_RETURNED,
               ROUND(COALESCE(SUM(r.RETURN_QUANTITY),0)
                     /NULLIF(SUM(f.ORDER_QUANTITY),0)*100,2)    AS RETURN_RATE_PCT
        FROM FACT_SALES f JOIN DIM_PRODUCT p ON f.PRODUCT_KEY=p.PRODUCT_KEY
        LEFT JOIN FACT_RETURNS r ON f.PRODUCT_KEY=r.PRODUCT_KEY
        GROUP BY 1,2 HAVING SUM(f.ORDER_QUANTITY)>0
        ORDER BY RETURN_RATE_PCT DESC LIMIT 15""")

    conn.close()
    return d


def load_from_csv():
    print("  Gold layer not available - reading local CSVs")
    sales = pd.concat([
        pd.read_csv(os.path.join(CSV, f"Sales Data {y}.csv"), encoding="latin1")
        for y in [2020, 2021, 2022]], ignore_index=True)
    products  = pd.read_csv(os.path.join(CSV,"Product Lookup.csv"),            encoding="latin1")
    cats      = pd.read_csv(os.path.join(CSV,"Product Categories Lookup.csv"), encoding="latin1")
    subcats   = pd.read_csv(os.path.join(CSV,"Subcategories Lookup.csv"),      encoding="latin1")
    territory = pd.read_csv(os.path.join(CSV,"Territory Lookup.csv"),          encoding="latin1")
    customers = pd.read_csv(os.path.join(CSV,"Customer Lookup.csv"),           encoding="latin1")
    returns   = pd.read_csv(os.path.join(CSV,"Returns Data.csv"),              encoding="latin1")

    products = products.merge(subcats, on="ProductSubcategoryKey", how="left")
    products = products.merge(cats,    on="ProductCategoryKey",    how="left")

    for df_, col in [(products,"ProductKey"),(customers,"CustomerKey"),
                     (territory,"SalesTerritoryKey")]:
        df_[col] = pd.to_numeric(df_[col], errors="coerce")
    for col in ["ProductKey","CustomerKey","TerritoryKey"]:
        sales[col] = pd.to_numeric(sales[col], errors="coerce")

    sales = (sales
             .merge(products,  on="ProductKey",  how="left")
             .merge(territory, left_on="TerritoryKey", right_on="SalesTerritoryKey", how="left")
             .merge(customers, on="CustomerKey",  how="left"))

    sales["OrderDate"] = pd.to_datetime(sales["OrderDate"], errors="coerce")
    sales["Year"]  = sales["OrderDate"].dt.year
    sales["Month"] = sales["OrderDate"].dt.month
    sales["Revenue"] = (sales["OrderQuantity"] * sales["ProductPrice"]).round(2)
    sales["Cost"]    = (sales["OrderQuantity"] * sales["ProductCost"]).round(2)
    sales["Profit"]  = (sales["Revenue"] - sales["Cost"]).round(2)
    sales["Margin"]  = (sales["Profit"] / sales["Revenue"].replace(0, pd.NA) * 100).round(2)

    def _tier(p):
        if p < 50:   return "Budget (<$50)"
        if p < 200:  return "Mid-Range ($50-199)"
        if p < 500:  return "Premium ($200-499)"
        return "Luxury ($500+)"
    sales["PriceTier"] = sales["ProductPrice"].apply(_tier)

    MN = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",
          7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}
    d = {}

    sy = (sales.groupby("Year")
               .agg(TOTAL_ORDERS=("OrderNumber","nunique"),
                    UNITS_SOLD=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"),
                    PROFIT=("Profit","sum"))
               .reset_index().rename(columns={"Year":"YEAR"}))
    sy[["REVENUE","PROFIT"]] = sy[["REVENUE","PROFIT"]].round(2)
    sy["MARGIN_PCT"]     = (sy["PROFIT"] / sy["REVENUE"] * 100).round(2)
    sy["YOY_GROWTH_PCT"] = sy["REVENUE"].pct_change().multiply(100).round(2)
    d["sales_by_year"] = sy

    mt = (sales.groupby(["Year","Month"])["Revenue"].sum().reset_index()
               .rename(columns={"Year":"YEAR","Month":"MONTH_NUM","Revenue":"REVENUE"}))
    mt["MONTH_NAME"] = mt["MONTH_NUM"].map(MN)
    mt["REVENUE"]    = mt["REVENUE"].round(2)
    d["monthly_trend"] = mt[["YEAR","MONTH_NUM","MONTH_NAME","REVENUE"]]

    cr = (sales.groupby("CategoryName")
               .agg(UNITS_SOLD=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"),
                    PROFIT=("Profit","sum"),
                    MARGIN_PCT=("Margin","mean"))
               .reset_index().rename(columns={"CategoryName":"CATEGORY_NAME"})
               .sort_values("REVENUE", ascending=False))
    cr[["REVENUE","PROFIT","MARGIN_PCT"]] = cr[["REVENUE","PROFIT","MARGIN_PCT"]].round(2)
    d["category_revenue"] = cr

    tp = (sales.groupby(["ProductName","CategoryName"])
               .agg(UNITS_SOLD=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"),
                    PROFIT=("Profit","sum"))
               .reset_index()
               .rename(columns={"ProductName":"PRODUCT_NAME","CategoryName":"CATEGORY_NAME"})
               .sort_values("REVENUE", ascending=False).head(10))
    tp[["REVENUE","PROFIT"]] = tp[["REVENUE","PROFIT"]].round(2)
    d["top_products"] = tp

    terr = (sales.groupby(["Country","Region"])
                 .agg(ORDERS=("OrderNumber","nunique"), REVENUE=("Revenue","sum"))
                 .reset_index()
                 .rename(columns={"Country":"COUNTRY","Region":"REGION"})
                 .sort_values("REVENUE", ascending=False))
    terr["REVENUE_PCT"] = (terr["REVENUE"] / terr["REVENUE"].sum() * 100).round(2)
    terr["REVENUE"]     = terr["REVENUE"].round(2)
    d["territory"] = terr

    gs = (sales.groupby("Gender")
               .agg(CUSTOMERS=("CustomerKey","nunique"),
                    ORDERS=("OrderNumber","nunique"),
                    REVENUE=("Revenue","sum"))
               .reset_index().rename(columns={"Gender":"GENDER"}))
    gs["REVENUE"] = gs["REVENUE"].round(2)
    d["gender_split"] = gs

    occ = (sales.groupby("Occupation")
                .agg(CUSTOMERS=("CustomerKey","nunique"),
                     ORDERS=("OrderNumber","nunique"),
                     REVENUE=("Revenue","sum"))
                .reset_index().rename(columns={"Occupation":"OCCUPATION"})
                .sort_values("REVENUE", ascending=False))
    occ["REVENUE"] = occ["REVENUE"].round(2)
    d["occupation_revenue"] = occ

    def _income_band(inc):
        if   inc < 30000:  return "Low (<$30K)"
        elif inc < 60000:  return "Average ($30-60K)"
        elif inc < 100000: return "High ($60-100K)"
        else:              return "Very High (>$100K)"

    sales["IncomeBand"] = sales["AnnualIncome"].apply(_income_band)
    cd = (sales.groupby(["IncomeBand","Occupation","Gender","MaritalStatus"])
               .agg(CUSTOMERS=("CustomerKey","nunique"),
                    ORDERS=("OrderNumber","nunique"),
                    REVENUE=("Revenue","sum"))
               .reset_index()
               .rename(columns={"IncomeBand":"INCOME_BAND","Occupation":"OCCUPATION",
                                 "Gender":"GENDER","MaritalStatus":"MARITAL_STATUS"})
               .sort_values("REVENUE", ascending=False))
    cd["REVENUE_PER_CUSTOMER"] = (cd["REVENUE"] / cd["CUSTOMERS"]).round(2)
    cd["REVENUE"] = cd["REVENUE"].round(2)
    d["customers_detail"] = cd

    rets     = returns.merge(products[["ProductKey","CategoryName"]], on="ProductKey", how="left")
    ret_cat  = rets.groupby("CategoryName")["ReturnQuantity"].sum().reset_index()
    sold_cat = sales.groupby("CategoryName")["OrderQuantity"].sum().reset_index()
    rc = (sold_cat.merge(ret_cat, on="CategoryName", how="left")
                  .rename(columns={"CategoryName":"CATEGORY_NAME",
                                   "OrderQuantity":"UNITS_SOLD",
                                   "ReturnQuantity":"UNITS_RETURNED"}))
    rc["UNITS_RETURNED"]  = rc["UNITS_RETURNED"].fillna(0).astype(int)
    rc["RETURN_RATE_PCT"] = (rc["UNITS_RETURNED"] / rc["UNITS_SOLD"] * 100).round(2)
    d["returns_by_category"] = rc.sort_values("RETURN_RATE_PCT", ascending=False)

    units_by_prod = (sales.groupby("ProductKey")["OrderQuantity"].sum()
                          .reset_index().rename(columns={"OrderQuantity":"UNITS_SOLD"}))
    ret_by_prod   = (returns.groupby("ProductKey")["ReturnQuantity"].sum()
                            .reset_index().rename(columns={"ReturnQuantity":"UNITS_RETURNED"}))
    rd = (units_by_prod
          .merge(ret_by_prod, on="ProductKey", how="left")
          .merge(products[["ProductKey","ProductName","CategoryName"]], on="ProductKey", how="left"))
    rd["UNITS_RETURNED"]  = rd["UNITS_RETURNED"].fillna(0).astype(int)
    rd["RETURN_RATE_PCT"] = (rd["UNITS_RETURNED"] / rd["UNITS_SOLD"] * 100).round(2)
    d["returns_detail"] = (rd.rename(columns={"ProductName":"PRODUCT_NAME","CategoryName":"CATEGORY_NAME"})
                             [["PRODUCT_NAME","CATEGORY_NAME","UNITS_SOLD","UNITS_RETURNED","RETURN_RATE_PCT"]]
                             .sort_values("RETURN_RATE_PCT", ascending=False)
                             .head(15))

    pt =(sales.groupby("PriceTier")
               .agg(ORDERS=("OrderNumber","nunique"),
                    UNITS_SOLD=("OrderQuantity","sum"),
                    REVENUE=("Revenue","sum"))
               .reset_index().rename(columns={"PriceTier":"PRICE_TIER"})
               .sort_values("REVENUE", ascending=False))
    pt["REVENUE_PCT"] = (pt["REVENUE"] / pt["REVENUE"].sum() * 100).round(2)
    pt["REVENUE"]     = pt["REVENUE"].round(2)
    d["price_tier"] = pt

    return d


def get_data():
    try:
        return load_from_snowflake()
    except Exception as e:
        print(f"  Snowflake unavailable ({e.__class__.__name__}) - using CSV fallback")
        return load_from_csv()


def compute_kpis(d):
    sy  = d["sales_by_year"]
    rev = sy["REVENUE"].sum()
    pro = sy["PROFIT"].sum()
    ord_ = sy["TOTAL_ORDERS"].sum()
    return {
        "total_revenue":   rev,
        "total_profit":    pro,
        "total_orders":    int(ord_),
        "avg_order_value": round(rev / ord_, 2) if ord_ > 0 else 0,
        "overall_margin":  round(pro / rev * 100, 1) if rev > 0 else 0,
        "top_country":     d["territory"].iloc[0]["COUNTRY"]           if len(d["territory"])         > 0 else "N/A",
        "top_category":    d["category_revenue"].iloc[0]["CATEGORY_NAME"] if len(d["category_revenue"]) > 0 else "N/A",
        "top_product":     d["top_products"].iloc[0]["PRODUCT_NAME"]   if len(d["top_products"])      > 0 else "N/A",
        "yoy_2021":        float(sy.iloc[1]["YOY_GROWTH_PCT"]) if len(sy) > 1 else None,
        "yoy_2022":        float(sy.iloc[2]["YOY_GROWTH_PCT"]) if len(sy) > 2 else None,
    }


# ══════════════════════════════════════════════════════════════════════════════
# 3.  CHART GENERATION (matplotlib -> BytesIO PNG)
# ══════════════════════════════════════════════════════════════════════════════

def _donut_pie(ax, values, labels, colors, title):
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, colors=colors,
        autopct="%1.1f%%", startangle=90,
        pctdistance=0.78, wedgeprops=dict(width=0.5, edgecolor="white", linewidth=1.5))
    for t in texts:     t.set_fontsize(8)
    for t in autotexts: t.set_fontsize(8); t.set_fontweight("bold"); t.set_color("white")
    ax.set_title(title, pad=12)


def chart_revenue_by_year(df):
    """Grouped bar: Revenue vs Profit per year with YoY growth labels."""
    fig, ax = plt.subplots(figsize=(7, 4))
    years = df["YEAR"].astype(str).tolist()
    x, w  = list(range(len(years))), 0.35
    b1 = ax.bar([i - w/2 for i in x], df["REVENUE"], w, label="Revenue", color=PALETTE[0])
    b2 = ax.bar([i + w/2 for i in x], df["PROFIT"],  w, label="Profit",  color=PALETTE[4])
    ax.set_xticks(x); ax.set_xticklabels(years)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Annual Revenue vs Profit"); ax.set_xlabel("Year"); ax.set_ylabel("USD")
    ax.legend(framealpha=0.8)
    for bar in b1:
        h = bar.get_height()
        ax.text(bar.get_x()+bar.get_width()/2, h*1.01, _fmt(h), ha="center", va="bottom", fontsize=8, fontweight="bold")
    for i, row in enumerate(df.itertuples()):
        yoy = getattr(row, "YOY_GROWTH_PCT", None)
        if pd.notna(yoy) and i > 0:
            c = "#27AE60" if yoy >= 0 else "#922B21"
            ax.annotate(f"YoY: {yoy:+.1f}%", xy=(i-w/2, row.REVENUE),
                        xytext=(0, 22), textcoords="offset points",
                        ha="center", fontsize=8, color=c, fontweight="bold")
    fig.tight_layout()
    return _savefig(fig)


def chart_monthly_trend(df):
    """Line chart: monthly revenue trend per year (one line per year)."""
    fig, ax = plt.subplots(figsize=(9, 4))
    pivot   = df.pivot(index="MONTH_NUM", columns="YEAR", values="REVENUE")
    labels  = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    clrs    = [PALETTE[0], PALETTE[1], PALETTE[5]]
    for i, (yr, col) in enumerate(pivot.items()):
        ax.plot(pivot.index, col.values, marker="o", linewidth=2,
                markersize=5, label=str(int(yr)), color=clrs[i % len(clrs)])
    ax.set_xticks(range(1, 13)); ax.set_xticklabels(labels, rotation=30, ha="right")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Monthly Revenue Trend (2020-2022)")
    ax.set_xlabel("Month"); ax.set_ylabel("Revenue (USD)")
    ax.legend(title="Year", framealpha=0.8)
    fig.tight_layout()
    return _savefig(fig)


def chart_category_bar(df):
    """Horizontal bar: revenue by product category."""
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    clrs = PALETTE[:len(df)]
    bars = ax.barh(df["CATEGORY_NAME"], df["REVENUE"], color=clrs, edgecolor="white")
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Revenue by Product Category"); ax.set_xlabel("Revenue (USD)")
    for bar in bars:
        w = bar.get_width()
        ax.text(w*1.01, bar.get_y()+bar.get_height()/2, _fmt(w), va="center", fontsize=8)
    ax.invert_yaxis(); fig.tight_layout()
    return _savefig(fig)


def chart_category_pie(df):
    """Donut pie: revenue share by category."""
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    _donut_pie(ax, df["REVENUE"], df["CATEGORY_NAME"], PALETTE[:len(df)], "Revenue Share by Category")
    fig.tight_layout()
    return _savefig(fig)


def chart_top_products_bar(df):
    """Horizontal bar: top 10 products by revenue."""
    fig, ax = plt.subplots(figsize=(9, 5))
    short = [n[:30] + ".." if len(n) > 32 else n for n in df["PRODUCT_NAME"]]
    bars  = ax.barh(short, df["REVENUE"], color=PALETTE[1], edgecolor="white")
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Top 10 Products by Revenue"); ax.set_xlabel("Revenue (USD)")
    for bar in bars:
        w = bar.get_width()
        ax.text(w*1.01, bar.get_y()+bar.get_height()/2, _fmt(w), va="center", fontsize=8)
    ax.invert_yaxis(); fig.tight_layout()
    return _savefig(fig)


def chart_territory_bar(df):
    """Horizontal bar: revenue by country."""
    cdf = df.groupby("COUNTRY")["REVENUE"].sum().reset_index().sort_values("REVENUE", ascending=False)
    fig, ax = plt.subplots(figsize=(6.5, 4))
    bars = ax.barh(cdf["COUNTRY"], cdf["REVENUE"], color=PALETTE[6], edgecolor="white")
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Revenue by Country"); ax.set_xlabel("Revenue (USD)")
    for bar in bars:
        w = bar.get_width()
        ax.text(w*1.01, bar.get_y()+bar.get_height()/2, _fmt(w), va="center", fontsize=8)
    ax.invert_yaxis(); fig.tight_layout()
    return _savefig(fig)


def chart_territory_pie(df):
    """Donut pie: revenue share by country."""
    cdf = df.groupby("COUNTRY")["REVENUE"].sum().reset_index().sort_values("REVENUE", ascending=False)
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    _donut_pie(ax, cdf["REVENUE"], cdf["COUNTRY"], PALETTE[4:4+len(cdf)], "Revenue Share by Country")
    fig.tight_layout()
    return _savefig(fig)


def chart_occupation_bar(df):
    """Vertical bar: revenue by customer occupation."""
    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(df["OCCUPATION"], df["REVENUE"], color=PALETTE[7], edgecolor="white")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Revenue by Customer Occupation")
    ax.set_xlabel("Occupation"); ax.set_ylabel("Revenue (USD)")
    plt.xticks(rotation=20, ha="right")
    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x()+bar.get_width()/2, h*1.01, _fmt(h), ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    return _savefig(fig)


def chart_gender_pie(df):
    """Donut pie: revenue split by gender."""
    fig, ax = plt.subplots(figsize=(5, 4.5))
    _donut_pie(ax, df["REVENUE"], df["GENDER"], [PALETTE[0], PALETTE[4]], "Revenue by Gender")
    fig.tight_layout()
    return _savefig(fig)


def chart_returns_bar(df):
    """Horizontal bar: return rate by product category."""
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    bars = ax.barh(df["CATEGORY_NAME"], df["RETURN_RATE_PCT"], color=PALETTE[10], edgecolor="white")
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"{v:.1f}%"))
    ax.set_title("Return Rate by Product Category"); ax.set_xlabel("Return Rate (%)")
    for bar in bars:
        w = bar.get_width()
        ax.text(w*1.01, bar.get_y()+bar.get_height()/2, f"{w:.1f}%", va="center", fontsize=8)
    ax.invert_yaxis(); fig.tight_layout()
    return _savefig(fig)


def chart_price_tier_pie(df):
    """Donut pie: revenue share by price tier."""
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    _donut_pie(ax, df["REVENUE"], df["PRICE_TIER"], PALETTE[2:2+len(df)], "Revenue by Price Tier")
    fig.tight_layout()
    return _savefig(fig)


def chart_price_tier_bar(df):
    """Vertical bar: revenue by price tier."""
    fig, ax = plt.subplots(figsize=(6.5, 4))
    bars = ax.bar(df["PRICE_TIER"], df["REVENUE"], color=PALETTE[2:2+len(df)], edgecolor="white")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: _fmt(v)))
    ax.set_title("Revenue by Price Tier")
    ax.set_xlabel("Price Tier"); ax.set_ylabel("Revenue (USD)")
    plt.xticks(rotation=10, ha="right")
    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x()+bar.get_width()/2, h*1.01, _fmt(h), ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    return _savefig(fig)


def make_all_charts(d):
    print("  Generating 12 charts...")
    return {
        "revenue_by_year": chart_revenue_by_year(d["sales_by_year"]),
        "monthly_trend":   chart_monthly_trend(d["monthly_trend"]),
        "category_bar":    chart_category_bar(d["category_revenue"]),
        "category_pie":    chart_category_pie(d["category_revenue"]),
        "top_products":    chart_top_products_bar(d["top_products"]),
        "territory_bar":   chart_territory_bar(d["territory"]),
        "territory_pie":   chart_territory_pie(d["territory"]),
        "occupation_bar":  chart_occupation_bar(d["occupation_revenue"]),
        "gender_pie":      chart_gender_pie(d["gender_split"]),
        "returns_bar":     chart_returns_bar(d["returns_by_category"]),
        "price_tier_pie":  chart_price_tier_pie(d["price_tier"]),
        "price_tier_bar":  chart_price_tier_bar(d["price_tier"]),
    }


# ══════════════════════════════════════════════════════════════════════════════
# 4.  EXCEL REPORT
# ══════════════════════════════════════════════════════════════════════════════

def _xl_table(ws, df, start_row=1, hdr_hex="1F4E79"):
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HDR  = PatternFill("solid", fgColor=hdr_hex)
    ALT  = PatternFill("solid", fgColor="D6E4F0")
    WFNT = Font(bold=True, color="FFFFFF", size=9)
    NFNT = Font(size=9)
    CTR  = Alignment(horizontal="center", vertical="center")
    thin = Side(style="thin", color="CCCCCC")
    BDR  = Border(left=thin, right=thin, top=thin, bottom=thin)

    for ci, col in enumerate(df.columns, 1):
        c = ws.cell(start_row, ci, col.replace("_", " "))
        c.fill = HDR; c.font = WFNT; c.alignment = CTR; c.border = BDR
        ws.row_dimensions[start_row].height = 18

    for ri, row in enumerate(df.itertuples(index=False), start_row + 1):
        fill = ALT if (ri - start_row) % 2 == 0 else None
        for ci, val in enumerate(row, 1):
            c = ws.cell(ri, ci, val)
            c.font = NFNT; c.alignment = CTR; c.border = BDR
            if fill: c.fill = fill

    for ci in range(1, len(df.columns) + 1):
        col_vals = [str(df.columns[ci-1])] + [str(v) for v in df.iloc[:, ci-1]]
        width    = min(max(len(v) for v in col_vals) + 3, 30)
        ws.column_dimensions[get_column_letter(ci)].width = width


def _xl_image(ws, chart_bytes, anchor, w_px=520, h_px=300):
    from openpyxl.drawing.image import Image as XLImage
    img = XLImage(io.BytesIO(chart_bytes))
    img.width = w_px; img.height = h_px
    ws.add_image(img, anchor)


def _xl_title(ws, text, row, ncols, hex_color="1F4E79"):
    from openpyxl.styles import Font, Alignment, PatternFill
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    c = ws.cell(row, 1, text)
    c.font      = Font(bold=True, size=13, color="FFFFFF")
    c.fill      = PatternFill("solid", fgColor=hex_color)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row].height = 22


def build_excel(d, charts, kpis):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ── Cover / KPI Dashboard ─────────────────────────────────────────────────
    cov = wb.create_sheet("KPI Dashboard")
    cov.sheet_view.showGridLines = False
    cov.column_dimensions["A"].width = 3
    for col in ["B","C","D","E"]: cov.column_dimensions[col].width = 20

    def kcell(r, c, v, bold=False, sz=11, clr="000000", bg=None):
        cell = cov.cell(r, c, v)
        cell.font      = Font(bold=bold, size=sz, color=clr)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        if bg: cell.fill = PatternFill("solid", fgColor=bg)
        return cell

    cov.row_dimensions[2].height = 36
    cov.merge_cells("B2:E2")
    kcell(2, 2, "Adventure Works Sales Report", bold=True, sz=20, clr="1F4E79")
    cov.row_dimensions[3].height = 16
    cov.merge_cells("B3:E3")
    kcell(3, 2, f"Generated: {TODAY}  |  Source: Snowflake ADVENTURE_WORKS_DB.GOLD", sz=10, clr="666666")

    tiles = [
        ("Total Revenue",    _fmt(kpis["total_revenue"]),  "1F4E79"),
        ("Total Profit",     _fmt(kpis["total_profit"]),   "145A32"),
        ("Total Orders",     f"{kpis['total_orders']:,}",  "6C3483"),
        ("Avg Order Value",  _fmt(kpis["avg_order_value"]), "784212"),
        ("Overall Margin",   f"{kpis['overall_margin']:.1f}%", "922B21"),
        ("Top Country",      kpis["top_country"],           "1F4E79"),
        ("Top Category",     kpis["top_category"],          "145A32"),
        ("YoY Growth 2022",
         f"{kpis['yoy_2022']:+.1f}%" if kpis.get("yoy_2022") is not None else "N/A",
         "27AE60" if (kpis.get("yoy_2022") or 0) >= 0 else "922B21"),
    ]
    for i, (lbl, val, bg) in enumerate(tiles):
        row = 5 + (i // 4) * 4
        col = 2 + (i % 4)
        cov.row_dimensions[row].height   = 18
        cov.row_dimensions[row+1].height = 26
        cov.row_dimensions[row+2].height = 8
        kcell(row,   col, lbl, bold=True, sz=9,  clr="666666")
        kcell(row+1, col, val, bold=True, sz=14, clr="FFFFFF", bg=bg)

    # embed revenue_by_year chart on cover
    _xl_image(cov, charts["revenue_by_year"], "B13", w_px=560, h_px=300)

    # ── Sheet 2: Sales Trend ──────────────────────────────────────────────────
    ws2 = wb.create_sheet("Sales Trend")
    ws2.sheet_view.showGridLines = False
    _xl_title(ws2, "Annual Sales Summary", 1, 7)
    _xl_table(ws2, d["sales_by_year"], start_row=2, hdr_hex="1F4E79")
    _xl_image(ws2, charts["monthly_trend"], "A10", w_px=700, h_px=300)

    # ── Sheet 3: Product Analysis ─────────────────────────────────────────────
    ws3 = wb.create_sheet("Product Analysis")
    ws3.sheet_view.showGridLines = False
    _xl_title(ws3, "Revenue by Category", 1, 5)
    _xl_table(ws3, d["category_revenue"], start_row=2, hdr_hex="145A32")
    _xl_image(ws3, charts["category_bar"], "A9",  w_px=520, h_px=260)
    _xl_image(ws3, charts["category_pie"], "J9",  w_px=440, h_px=340)

    row_offset = 9 + 20  # approx rows for chart + gap
    _xl_title(ws3, "Top 10 Products by Revenue", row_offset, 5, hex_color="145A32")
    _xl_table(ws3, d["top_products"], start_row=row_offset + 1, hdr_hex="145A32")
    _xl_image(ws3, charts["top_products"], f"A{row_offset + 14}", w_px=700, h_px=340)

    # ── Sheet 4: Territory ────────────────────────────────────────────────────
    ws4 = wb.create_sheet("Territory Analysis")
    ws4.sheet_view.showGridLines = False
    _xl_title(ws4, "Revenue by Territory", 1, 5)
    _xl_table(ws4, d["territory"], start_row=2, hdr_hex="6C3483")
    _xl_image(ws4, charts["territory_bar"], "A14", w_px=520, h_px=280)
    _xl_image(ws4, charts["territory_pie"], "J14", w_px=440, h_px=340)

    # ── Sheet 5: Customer Analysis ────────────────────────────────────────────
    ws5 = wb.create_sheet("Customer Analysis")
    ws5.sheet_view.showGridLines = False
    _xl_title(ws5, "Revenue by Occupation", 1, 4)
    _xl_table(ws5, d["occupation_revenue"], start_row=2, hdr_hex="784212")
    _xl_image(ws5, charts["occupation_bar"], "A10", w_px=540, h_px=280)
    _xl_image(ws5, charts["gender_pie"],     "J10", w_px=400, h_px=320)

    _xl_title(ws5, "Gender Split", 26, 4, hex_color="784212")
    _xl_table(ws5, d["gender_split"], start_row=27, hdr_hex="784212")

    _xl_title(ws5, "Customer Segments — Full Breakdown (Income / Occupation / Gender / Marital Status)",
              32, 8, hex_color="4A235A")
    _xl_table(ws5, d["customers_detail"], start_row=33, hdr_hex="4A235A")

    # ── Sheet 6: Returns & Price Tiers ────────────────────────────────────────
    ws6 = wb.create_sheet("Returns & Price Tiers")
    ws6.sheet_view.showGridLines = False
    _xl_title(ws6, "Returns by Category", 1, 4)
    _xl_table(ws6, d["returns_by_category"], start_row=2, hdr_hex="922B21")
    _xl_image(ws6, charts["returns_bar"], "A8",  w_px=520, h_px=260)

    _xl_title(ws6, "Returns by Product — Top 15 by Return Rate", 22, 5, hex_color="922B21")
    _xl_table(ws6, d["returns_detail"], start_row=23, hdr_hex="922B21")

    _xl_title(ws6, "Revenue by Price Tier", 42, 5, hex_color="784212")
    _xl_table(ws6, d["price_tier"], start_row=43, hdr_hex="784212")
    _xl_image(ws6, charts["price_tier_bar"], "A48", w_px=520, h_px=280)
    _xl_image(ws6, charts["price_tier_pie"], "J48", w_px=440, h_px=340)

    wb.save(XLSX)
    print(f"  Excel saved -> {XLSX}")


# ══════════════════════════════════════════════════════════════════════════════
# 5.  PDF REPORT
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(d, charts, kpis):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, PageBreak, Image as RLImage,
                                    HRFlowable)

    LA    = landscape(A4)
    doc   = SimpleDocTemplate(PDF, pagesize=LA,
                               leftMargin=1.5*cm, rightMargin=1.5*cm,
                               topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    BLUE   = colors.HexColor("#1F4E79")
    GREEN  = colors.HexColor("#145A32")
    PURPLE = colors.HexColor("#6C3483")
    BROWN  = colors.HexColor("#784212")
    RED    = colors.HexColor("#922B21")
    LBLUE  = colors.HexColor("#D6E4F0")
    WHITE  = colors.white
    LGREY  = colors.HexColor("#F5F5F5")

    H0 = ParagraphStyle("h0", parent=styles["Title"],   textColor=BLUE,  fontSize=22, spaceAfter=8)
    H1 = ParagraphStyle("h1", parent=styles["Heading1"],textColor=BLUE,  fontSize=15, spaceAfter=6, spaceBefore=10)
    H2 = ParagraphStyle("h2", parent=styles["Heading2"],textColor=GREEN, fontSize=12, spaceAfter=4)
    SM = ParagraphStyle("sm", parent=styles["Normal"],  fontSize=9,      spaceAfter=3, leading=13)
    INS= ParagraphStyle("ins",parent=styles["Normal"],  fontSize=9,      spaceAfter=3,
                        backColor=LGREY, borderPad=4, leading=13)

    PAGE_W = LA[0] - 3*cm

    def rl_img(key, width=PAGE_W, height=10*cm):
        return RLImage(io.BytesIO(charts[key]), width=width, height=height)

    def df_table(df, hdr_clr=BLUE):
        header = [c.replace("_"," ") for c in df.columns]
        rows   = [header] + [[str(v) for v in r] for r in df.itertuples(index=False)]
        cw     = PAGE_W / len(df.columns)
        t = Table(rows, colWidths=[cw]*len(df.columns), repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,0),  hdr_clr),
            ("TEXTCOLOR",   (0,0),(-1,0),  WHITE),
            ("FONTNAME",    (0,0),(-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0),(-1,-1), 7.5),
            ("ALIGN",       (0,0),(-1,-1), "CENTER"),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE, LBLUE]),
            ("GRID",        (0,0),(-1,-1), 0.4, colors.grey),
            ("TOPPADDING",  (0,0),(-1,-1), 3),
            ("BOTTOMPADDING",(0,0),(-1,-1), 3),
        ]))
        return t

    def kpi_row(items):
        cells = [[Paragraph(f"<b>{v}</b><br/><font size=7 color='#666666'>{l}</font>", SM)]
                 for l, v in items]
        t = Table([cells], colWidths=[PAGE_W/len(items)]*len(items))
        t.setStyle(TableStyle([
            ("BOX",         (0,0),(-1,-1), 0.5, BLUE),
            ("INNERGRID",   (0,0),(-1,-1), 0.5, LBLUE),
            ("BACKGROUND",  (0,0),(-1,-1), LGREY),
            ("ALIGN",       (0,0),(-1,-1), "CENTER"),
            ("VALIGN",      (0,0),(-1,-1), "MIDDLE"),
            ("TOPPADDING",  (0,0),(-1,-1), 6),
            ("BOTTOMPADDING",(0,0),(-1,-1), 6),
        ]))
        return t

    story = []

    # ── Cover page ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 2.5*cm))
    story.append(Paragraph("Adventure Works Sales Report", H0))
    story.append(HRFlowable(width=PAGE_W, color=BLUE, thickness=2))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(f"Generated: {TODAY}  |  "
                            "Source: Snowflake ADVENTURE_WORKS_DB.GOLD  |  "
                            "Account: uq57089.ap-southeast-7.aws", SM))
    story.append(Spacer(1, 0.6*cm))

    story.append(kpi_row([
        ("Total Revenue",   _fmt(kpis["total_revenue"])),
        ("Total Profit",    _fmt(kpis["total_profit"])),
        ("Total Orders",    f"{kpis['total_orders']:,}"),
        ("Avg Order Value", _fmt(kpis["avg_order_value"])),
        ("Overall Margin",  f"{kpis['overall_margin']:.1f}%"),
        ("YoY Growth 2022", f"{kpis['yoy_2022']:+.1f}%" if kpis.get("yoy_2022") is not None else "N/A"),
    ]))
    story.append(Spacer(1, 0.5*cm))
    story.append(rl_img("revenue_by_year", width=PAGE_W, height=10*cm))
    story.append(PageBreak())

    # ── Section 1: Sales Trend ────────────────────────────────────────────────
    story.append(Paragraph("1. Sales Performance", H1))
    sy  = d["sales_by_year"]
    best_yr = sy.loc[sy["REVENUE"].idxmax()]
    story.append(Paragraph(
        f"<b>Key Insight:</b> Three-year total revenue of {_fmt(kpis['total_revenue'])}. "
        f"Best year was <b>{int(best_yr['YEAR'])}</b> with revenue of "
        f"<b>{_fmt(best_yr['REVENUE'])}</b>."
        + (f" 2022 showed YoY growth of <b>{kpis['yoy_2022']:+.1f}%</b>."
           if kpis.get("yoy_2022") is not None else ""), SM))
    story.append(Spacer(1, 0.3*cm))
    story.append(df_table(sy, BLUE))
    story.append(Spacer(1, 0.4*cm))
    story.append(rl_img("monthly_trend", width=PAGE_W, height=9.5*cm))
    story.append(PageBreak())

    # ── Section 2: Product Analysis ───────────────────────────────────────────
    story.append(Paragraph("2. Product Analysis", H1))
    cr   = d["category_revenue"]
    top1 = cr.iloc[0]
    pct1 = round(top1["REVENUE"] / cr["REVENUE"].sum() * 100, 1)
    story.append(Paragraph(
        f"<b>Key Insight:</b> <b>{top1['CATEGORY_NAME']}</b> is the top category, "
        f"contributing <b>{_fmt(top1['REVENUE'])}</b> ({pct1}% of total). "
        f"Top product overall: <b>{kpis['top_product']}</b>.", SM))
    story.append(Spacer(1, 0.3*cm))

    half = PAGE_W / 2 - 0.3*cm
    img_row = Table([[rl_img("category_bar", width=half+1*cm, height=8.5*cm),
                      rl_img("category_pie", width=half-1*cm, height=8.5*cm)]],
                    colWidths=[half+1*cm, half-1*cm])
    story.append(img_row)
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph("Top 10 Products", H2))
    story.append(df_table(d["top_products"], GREEN))
    story.append(Spacer(1, 0.4*cm))
    story.append(rl_img("top_products", width=PAGE_W, height=9.5*cm))
    story.append(PageBreak())

    # ── Section 3: Territory Analysis ────────────────────────────────────────
    story.append(Paragraph("3. Territory Analysis", H1))
    terr = d["territory"]
    t1   = terr.iloc[0]
    story.append(Paragraph(
        f"<b>Key Insight:</b> <b>{t1['COUNTRY']}</b> leads with "
        f"{_fmt(t1['REVENUE'])} ({t1['REVENUE_PCT']:.1f}% revenue share). "
        f"Top 3 countries account for "
        f"{terr.head(3)['REVENUE_PCT'].sum():.1f}% of total revenue.", SM))
    story.append(Spacer(1, 0.3*cm))
    story.append(df_table(terr, PURPLE))
    story.append(Spacer(1, 0.4*cm))
    img_row2 = Table([[rl_img("territory_bar", width=half+1*cm, height=9*cm),
                       rl_img("territory_pie", width=half-1*cm, height=9*cm)]],
                     colWidths=[half+1*cm, half-1*cm])
    story.append(img_row2)
    story.append(PageBreak())

    # ── Section 4: Customer Analysis ─────────────────────────────────────────
    story.append(Paragraph("4. Customer Analysis", H1))
    occ  = d["occupation_revenue"]
    gs   = d["gender_split"]
    top_occ = occ.iloc[0]
    story.append(Paragraph(
        f"<b>Key Insight:</b> <b>{top_occ['OCCUPATION']}</b> customers generate the highest "
        f"revenue at {_fmt(top_occ['REVENUE'])} from {top_occ['CUSTOMERS']:,} customers. "
        f"Gender split (by revenue): "
        + ", ".join(f"{r['GENDER']} {r['REVENUE']/gs['REVENUE'].sum()*100:.0f}%"
                    for _, r in gs.iterrows()) + ".", SM))
    story.append(Spacer(1, 0.3*cm))
    story.append(df_table(occ, BROWN))
    story.append(Spacer(1, 0.4*cm))
    img_row3 = Table([[rl_img("occupation_bar", width=half+1.5*cm, height=9*cm),
                       rl_img("gender_pie",     width=half-1.5*cm, height=9*cm)]],
                     colWidths=[half+1.5*cm, half-1.5*cm])
    story.append(img_row3)
    story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("Customer Segments — Full Breakdown", H2))
    story.append(Paragraph(
        "Revenue segmented by income band, occupation, gender, and marital status. "
        "Shows customer count, order count, total revenue, and revenue per customer.", SM))
    story.append(Spacer(1, 0.3*cm))
    story.append(df_table(d["customers_detail"].head(20), PURPLE))
    story.append(PageBreak())

    # ── Section 5: Returns & Price Tiers ─────────────────────────────────────
    story.append(Paragraph("5. Returns & Price Tier Analysis", H1))
    rc   = d["returns_by_category"]
    pt   = d["price_tier"]
    top_ret = rc.iloc[0]
    story.append(Paragraph(
        f"<b>Returns Key Insight:</b> <b>{top_ret['CATEGORY_NAME']}</b> has the highest "
        f"return rate at <b>{top_ret['RETURN_RATE_PCT']:.1f}%</b> "
        f"({top_ret['UNITS_RETURNED']:,} units returned of {top_ret['UNITS_SOLD']:,} sold).", SM))
    story.append(Spacer(1, 0.2*cm))
    story.append(df_table(rc, RED))
    story.append(Spacer(1, 0.4*cm))
    story.append(rl_img("returns_bar", width=PAGE_W * 0.65, height=8*cm))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("Returns by Product — Top 15 by Return Rate", H2))
    story.append(Spacer(1, 0.2*cm))
    story.append(df_table(d["returns_detail"], RED))
    story.append(Spacer(1, 0.4*cm))

    top_tier = pt.iloc[0]
    story.append(Paragraph(
        f"<b>Price Tier Key Insight:</b> <b>{top_tier['PRICE_TIER']}</b> products "
        f"drive the most revenue at {_fmt(top_tier['REVENUE'])} "
        f"({top_tier['REVENUE_PCT']:.1f}% of total).", SM))
    story.append(Spacer(1, 0.2*cm))
    story.append(df_table(pt, RED))
    story.append(Spacer(1, 0.4*cm))
    img_row4 = Table([[rl_img("price_tier_bar", width=half+1*cm, height=9*cm),
                       rl_img("price_tier_pie", width=half-1*cm, height=9*cm)]],
                     colWidths=[half+1*cm, half-1*cm])
    story.append(img_row4)

    doc.build(story)
    print(f"  PDF  saved -> {PDF}")


# ══════════════════════════════════════════════════════════════════════════════
# 6.  WORD REPORT
# ══════════════════════════════════════════════════════════════════════════════

def build_word(d, charts, kpis):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE  = RGBColor(0x1F, 0x4E, 0x79)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GREEN = RGBColor(0x14, 0x5A, 0x32)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def _shd(cell, hex_colour):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_colour)
        tcPr.append(shd)

    def add_table(df, hdr_hex="1F4E79"):
        table = doc.add_table(rows=1+len(df), cols=len(df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(df.columns):
            c = table.rows[0].cells[i]
            c.text = col.replace("_", " ")
            _shd(c, hdr_hex)
            run = c.paragraphs[0].runs[0]
            run.bold = True; run.font.size = Pt(8); run.font.color.rgb = WHITE
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ri, row in enumerate(df.itertuples(index=False)):
            for ci, val in enumerate(row):
                c = table.rows[ri+1].cells[ci]
                c.text = str(val)
                if ri % 2 == 0: _shd(c, "D6E4F0")
                run = c.paragraphs[0].runs[0]
                run.font.size = Pt(8)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_w = Inches(6.2)
        cw = total_w / len(df.columns)
        for col in table.columns:
            for c in col.cells:
                c.width = cw
        doc.add_paragraph()

    def add_chart(key, width=Inches(5.8)):
        doc.add_picture(io.BytesIO(charts[key]), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = BLUE

    def subheading(text):
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.color.rgb = GREEN

    def insight(text):
        p = doc.add_paragraph()
        p.add_run("Key Insight: ").bold = True
        p.add_run(text)
        p.style.font.size = Pt(9)

    # ── Cover ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_heading("Adventure Works Sales Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs: run.font.color.rgb = BLUE; run.font.size = Pt(24)
    doc.add_paragraph()

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    for label, val in [("Generated", TODAY),
                       ("Source",    "Snowflake ADVENTURE_WORKS_DB.GOLD"),
                       ("Account",   "uq57089.ap-southeast-7.aws"),
                       ("Warehouse", "COMPUTE_WH")]:
        row = meta.rows[["Generated","Source","Account","Warehouse"].index(label)]
        row.cells[0].text = label; row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # KPI table
    kpi_data = [
        ["Total Revenue", "Total Profit", "Total Orders", "Avg Order Value", "Overall Margin", "YoY 2022"],
        [_fmt(kpis["total_revenue"]),
         _fmt(kpis["total_profit"]),
         f"{kpis['total_orders']:,}",
         _fmt(kpis["avg_order_value"]),
         f"{kpis['overall_margin']:.1f}%",
         f"{kpis['yoy_2022']:+.1f}%" if kpis.get("yoy_2022") is not None else "N/A"],
    ]
    ktbl = doc.add_table(rows=2, cols=6)
    ktbl.style = "Table Grid"
    for ci, val in enumerate(kpi_data[0]):
        c = ktbl.rows[0].cells[ci]; c.text = val
        _shd(c, "1F4E79")
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ci, val in enumerate(kpi_data[1]):
        c = ktbl.rows[1].cells[ci]; c.text = val
        _shd(c, "D6E4F0")
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_chart("revenue_by_year")
    doc.add_page_break()

    # ── Section 1: Sales Trend ────────────────────────────────────────────────
    heading("1. Sales Performance")
    sy      = d["sales_by_year"]
    best_yr = sy.loc[sy["REVENUE"].idxmax()]
    insight(f"Three-year total revenue of {_fmt(kpis['total_revenue'])}. "
            f"Best year: {int(best_yr['YEAR'])} with {_fmt(best_yr['REVENUE'])} revenue."
            + (f" 2022 YoY growth: {kpis['yoy_2022']:+.1f}%." if kpis.get("yoy_2022") is not None else ""))
    add_table(sy, "1F4E79")
    add_chart("monthly_trend")
    doc.add_page_break()

    # ── Section 2: Product Analysis ───────────────────────────────────────────
    heading("2. Product Analysis")
    cr   = d["category_revenue"]
    top1 = cr.iloc[0]
    pct1 = round(top1["REVENUE"] / cr["REVENUE"].sum() * 100, 1)
    insight(f"{top1['CATEGORY_NAME']} leads with {_fmt(top1['REVENUE'])} ({pct1}% of total). "
            f"Top product: {kpis['top_product']}.")
    subheading("Revenue by Category")
    add_table(cr, "145A32")
    add_chart("category_bar",  width=Inches(5.5))
    add_chart("category_pie",  width=Inches(4.5))
    subheading("Top 10 Products")
    add_table(d["top_products"], "145A32")
    add_chart("top_products",  width=Inches(5.8))
    doc.add_page_break()

    # ── Section 3: Territory Analysis ────────────────────────────────────────
    heading("3. Territory Analysis")
    terr = d["territory"]
    t1   = terr.iloc[0]
    insight(f"{t1['COUNTRY']} leads with {_fmt(t1['REVENUE'])} ({t1['REVENUE_PCT']:.1f}% share). "
            f"Top 3 countries: {terr.head(3)['REVENUE_PCT'].sum():.1f}% of total revenue.")
    add_table(terr, "6C3483")
    add_chart("territory_bar", width=Inches(5.5))
    add_chart("territory_pie", width=Inches(4.5))
    doc.add_page_break()

    # ── Section 4: Customer Analysis ─────────────────────────────────────────
    heading("4. Customer Analysis")
    occ     = d["occupation_revenue"]
    gs      = d["gender_split"]
    top_occ = occ.iloc[0]
    insight(f"{top_occ['OCCUPATION']} customers generate {_fmt(top_occ['REVENUE'])} "
            f"from {top_occ['CUSTOMERS']:,} customers. "
            + "Gender revenue split: "
            + ", ".join(f"{r['GENDER']} {r['REVENUE']/gs['REVENUE'].sum()*100:.0f}%"
                        for _, r in gs.iterrows()) + ".")
    add_table(occ, "784212")
    add_chart("occupation_bar", width=Inches(5.5))
    add_chart("gender_pie",     width=Inches(4.5))
    subheading("Customer Segments — Full Breakdown (Income / Occupation / Gender / Marital Status)")
    insight("Revenue segmented across all four demographic dimensions. "
            "Shows customer count, order count, total revenue, and revenue per customer.")
    add_table(d["customers_detail"].head(20), "4A235A")
    doc.add_page_break()

    # ── Section 5: Returns & Price Tiers ─────────────────────────────────────
    heading("5. Returns & Price Tier Analysis")
    rc      = d["returns_by_category"]
    pt      = d["price_tier"]
    top_ret = rc.iloc[0]
    top_tier= pt.iloc[0]
    insight(f"Returns: {top_ret['CATEGORY_NAME']} has highest return rate "
            f"{top_ret['RETURN_RATE_PCT']:.1f}% "
            f"({top_ret['UNITS_RETURNED']:,} of {top_ret['UNITS_SOLD']:,} units). "
            f"Price Tier: {top_tier['PRICE_TIER']} drives the most revenue "
            f"at {_fmt(top_tier['REVENUE'])} ({top_tier['REVENUE_PCT']:.1f}%).")
    subheading("Returns by Category")
    add_table(rc, "922B21")
    add_chart("returns_bar", width=Inches(5.5))
    subheading("Returns by Product — Top 15 by Return Rate")
    add_table(d["returns_detail"], "922B21")
    subheading("Revenue by Price Tier")
    add_table(pt, "922B21")
    add_chart("price_tier_bar", width=Inches(5.5))
    add_chart("price_tier_pie", width=Inches(4.5))

    doc.save(DOCX)
    print(f"  Word saved -> {DOCX}")


# ══════════════════════════════════════════════════════════════════════════════
# 7.  MAIN
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    print("Adventure Works Enhanced Report Generator v2")
    print("=" * 46)

    print("\n[1/5] Loading data...")
    d = get_data()
    kpis = compute_kpis(d)
    print(f"      {len(d['sales_by_year'])} years | "
          f"{len(d['top_products'])} products | "
          f"{len(d['territory'])} territories | "
          f"Revenue: {_fmt(kpis['total_revenue'])}")

    print("\n[2/5] Generating charts...")
    charts = make_all_charts(d)

    print("\n[3/5] Building Excel report...")
    build_excel(d, charts, kpis)

    print("\n[4/5] Building PDF report...")
    build_pdf(d, charts, kpis)

    print("\n[5/5] Building Word report...")
    build_word(d, charts, kpis)

    print("\nDone!")
    print(f"  Excel -> {XLSX}")
    print(f"  PDF   -> {PDF}")
    print(f"  Word  -> {DOCX}")
